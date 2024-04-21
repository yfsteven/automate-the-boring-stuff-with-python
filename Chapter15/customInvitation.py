import docx
from pathlib import Path

p = open(Path.cwd() / 'guests.txt')
lines = p.readlines()
for line in lines:
    doc = docx.Document()
    doc.styles['Normal'].font.name = 'Bush Script Std'
    doc.styles['Normal'].font.size = docx.shared.Pt(20)
    paragraph_format = doc.styles['Normal'].paragraph_format
    paragraph_format.line_spacing = 1
    paraObj1 = doc.add_paragraph('It would be a pleasure to have the company of')
    doc.paragraphs[0].runs[0].add_break()
    paraObj2 = doc.add_paragraph(line.strip())
    doc.paragraphs[1].runs[0].add_break()
    doc.paragraphs[1].runs[0].bold = True
    paraObj3 = doc.add_paragraph('at 1101 Memory Lane on the Evening of')
    doc.paragraphs[2].runs[0].add_break()
    paraObj4 = doc.add_paragraph('April 1st')
    doc.paragraphs[3].runs[0].add_break()
    doc.paragraphs[3].runs[0].italic = True
    paraObj5 = doc.add_paragraph("at 7 o'clock")
    doc.paragraphs[0].alignment = 1
    doc.paragraphs[1].alignment = 1
    doc.paragraphs[2].alignment = 1
    doc.paragraphs[3].alignment = 1
    doc.paragraphs[4].alignment = 1
    doc.save(line + 'invitation.docx')


